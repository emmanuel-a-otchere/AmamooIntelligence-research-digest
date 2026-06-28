#!/usr/bin/env python3
"""Generate docs/REQUIREMENTS.docx from docs/REQUIREMENTS.md using python-docx.

Since we don't have pandoc, this script parses the Markdown manually and applies
structured styles: headings, paragraphs, tables, code blocks, lists.

Goal: a Word document that any operator can read or share, formatted professionally.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parent.parent
SRC = REPO_ROOT / "docs" / "REQUIREMENTS.md"
DST = REPO_ROOT / "docs" / "REQUIREMENTS.docx"


# -------- Colour palette --------
COLOR_HEADING1 = RGBColor(0x1F, 0x3A, 0x5F)   # deep navy
COLOR_HEADING2 = RGBColor(0x2C, 0x5F, 0x82)   # mid navy
COLOR_HEADING3 = RGBColor(0x3D, 0x7A, 0xA5)   # lighter navy
COLOR_HEADING4 = RGBColor(0x55, 0x55, 0x55)   # dark grey
COLOR_CODE_BG  = RGBColor(0xF4, 0xF4, 0xF4)   # light grey
COLOR_TABLE_HEADER = RGBColor(0x1F, 0x3A, 0x5F)
COLOR_TABLE_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_TABLE_ALT = RGBColor(0xF7, 0xF9, 0xFB)
COLOR_LINK = RGBColor(0x06, 0x4B, 0x9E)
COLOR_QUOTE = RGBColor(0x55, 0x55, 0x55)
COLOR_QUOTE_BORDER = RGBColor(0xC8, 0xCE, 0xD7)


def shade_cell(cell, fill_hex: str) -> None:
    """Set the background fill colour of a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def add_horizontal_rule(doc) -> None:
    """Insert a horizontal rule (page break-like separator)."""
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BBBBBB")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_inline_runs(paragraph, text: str, base_size: Pt | None = None) -> None:
    """Parse inline Markdown (bold, italic, code, links) and add styled runs."""
    i = 0
    n = len(text)
    buf = ""
    while i < n:
        # Inline code: `...`
        if text[i] == "`":
            j = text.find("`", i + 1)
            if j == -1:
                buf += text[i]
                i += 1
                continue
            if buf:
                r = paragraph.add_run(buf)
                if base_size: r.font.size = base_size
                buf = ""
            code_text = text[i + 1:j]
            r = paragraph.add_run(code_text)
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
            rPr = r._element.get_or_add_rPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F0F0F0")
            rPr.append(shd)
            i = j + 1
            continue
        # Bold: **...**
        if text[i:i + 2] == "**":
            j = text.find("**", i + 2)
            if j == -1:
                buf += text[i]
                i += 1
                continue
            if buf:
                r = paragraph.add_run(buf)
                if base_size: r.font.size = base_size
                buf = ""
            inner = text[i + 2:j]
            r = paragraph.add_run(inner)
            r.bold = True
            if base_size: r.font.size = base_size
            i = j + 2
            continue
        # Italic: *...*
        if text[i] == "*" and (i == 0 or text[i - 1] != "*"):
            j = text.find("*", i + 1)
            if j == -1 or j == i + 1:
                buf += text[i]
                i += 1
                continue
            if buf:
                r = paragraph.add_run(buf)
                if base_size: r.font.size = base_size
                buf = ""
            inner = text[i + 1:j]
            r = paragraph.add_run(inner)
            r.italic = True
            if base_size: r.font.size = base_size
            i = j + 1
            continue
        # Link: [text](url)
        if text[i] == "[":
            close = text.find("]", i + 1)
            paren_open = text.find("(", close + 1) if close != -1 else -1
            paren_close = text.find(")", paren_open + 1) if paren_open != -1 else -1
            if close != -1 and paren_open == close + 1 and paren_close != -1:
                if buf:
                    r = paragraph.add_run(buf)
                    if base_size: r.font.size = base_size
                    buf = ""
                link_text = text[i + 1:close]
                link_url = text[paren_open + 1:paren_close]
                r = paragraph.add_run(link_text)
                r.font.color.rgb = COLOR_LINK
                r.font.underline = True
                if base_size: r.font.size = base_size
                # Add hyperlink relationship
                part = paragraph.part
                r_id = part.relate_to(
                    link_url,
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                    is_external=True,
                )
                hyperlink = OxmlElement("w:hyperlink")
                hyperlink.set(qn("r:id"), r_id)
                run = r._element
                hyperlink.append(run)
                paragraph._p.append(hyperlink)
                # Remove the original run from paragraph (it's now inside hyperlink)
                run.getparent().remove(run)
                i = paren_close + 1
                continue
        # Plain character
        buf += text[i]
        i += 1
    if buf:
        r = paragraph.add_run(buf)
        if base_size: r.font.size = base_size


def render_table(doc, header_cells: list[str], body_rows: list[list[str]]):
    """Render a Markdown pipe-table as a Word table with styled header."""
    if not header_cells:
        return
    table = doc.add_table(rows=1 + len(body_rows), cols=len(header_cells))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Header row
    hdr = table.rows[0]
    for i, txt in enumerate(header_cells):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        add_inline_runs(p, txt, base_size=Pt(9))
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = COLOR_TABLE_HEADER_TEXT
        shade_cell(cell, "1F3A5F")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Body rows
    for ri, row in enumerate(body_rows):
        for ci, txt in enumerate(row):
            cell = table.rows[1 + ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_inline_runs(p, txt, base_size=Pt(9))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if ri % 2 == 1:
                shade_cell(cell, "F7F9FB")


def render_code_block(doc, code: str) -> None:
    """Render a fenced code block as a single-cell grey-shaded table."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell = table.rows[0].cells[0]
    shade_cell(cell, "F4F4F4")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    for line in code.splitlines():
        if p.runs:
            p.add_run().add_break()
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def render_blockquote(doc, text: str) -> None:
    """Render a Markdown blockquote as an indented italic paragraph with a left border."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell = table.rows[0].cells[0]
    # Left border accent
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), "C8CED7")
    tc_borders.append(left)
    tc_pr.append(tc_borders)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    # Strip leading "> "
    for line in text.splitlines():
        clean = re.sub(r"^>\s?", "", line)
        if p.runs:
            p.add_run().add_break()
        add_inline_runs(p, clean if clean else " ", base_size=Pt(10.5))
        for run in p.runs:
            run.italic = True
            run.font.color.rgb = COLOR_QUOTE


def parse_table_block(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    """Parse a Markdown pipe-table into (header, rows). Assumes a header divider
    on the second line (---|---)."""
    def split_row(s: str) -> list[str]:
        s = s.strip().strip("|")
        return [c.strip() for c in s.split("|")]
    header = split_row(lines[0])
    rows = [split_row(line) for line in lines[2:]]
    return header, rows


def add_page_break(doc) -> None:
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(7)  # 7 = WD_BREAK.PAGE


# -------- Main conversion loop --------

def convert(md_text: str, doc: Document) -> None:
    lines = md_text.splitlines()
    i = 0
    n = len(lines)
    in_code_block = False
    code_buf: list[str] = []
    in_table = False
    table_buf: list[str] = []

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Code block
        if stripped.startswith("```"):
            if in_code_block:
                # Closing fence
                render_code_block(doc, "\n".join(code_buf))
                code_buf = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        if in_code_block:
            code_buf.append(line)
            i += 1
            continue

        # Table
        if "|" in line and i + 1 < n and re.match(r"^\s*\|[\s\-:|]+\|\s*$", lines[i + 1]):
            in_table = True
            table_buf = []
            while i < n and "|" in lines[i]:
                table_buf.append(lines[i])
                i += 1
            header, rows = parse_table_block(table_buf)
            render_table(doc, header, rows)
            in_table = False
            table_buf = []
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1:
                p = doc.add_heading(level=1)
            elif level == 2:
                p = doc.add_heading(level=2)
            elif level == 3:
                p = doc.add_heading(level=3)
            elif level == 4:
                p = doc.add_heading(level=4)
            elif level == 5:
                p = doc.add_heading(level=5)
            else:
                p = doc.add_heading(level=6)
            p.text = ""
            add_inline_runs(p, text)
            # Apply colours
            if level == 1:
                p.runs[0].font.color.rgb = COLOR_HEADING1
                p.runs[0].font.size = Pt(22)
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.keep_with_next = True
            elif level == 2:
                p.runs[0].font.color.rgb = COLOR_HEADING2
                p.runs[0].font.size = Pt(16)
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_with_next = True
            elif level == 3:
                p.runs[0].font.color.rgb = COLOR_HEADING3
                p.runs[0].font.size = Pt(13)
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.keep_with_next = True
            else:
                p.runs[0].font.color.rgb = COLOR_HEADING4
                p.runs[0].font.size = Pt(11)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.keep_with_next = True
            i += 1
            continue

        # Horizontal rule
        if stripped in ("---", "***", "___"):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            quote_buf = []
            while i < n and lines[i].strip().startswith(">"):
                quote_buf.append(lines[i])
                i += 1
            render_blockquote(doc, "\n".join(quote_buf))
            continue

        # Table of contents - skip the entire TOC block (Word can regen if user wants)
        if stripped in ("## Table of Contents", "## Contents", "## Table of contents", "## table of contents"):
            i += 1
            # Skip the indented TOC list items following the heading
            while i < n and (lines[i].strip().startswith("- [") or not lines[i].strip()):
                i += 1
            continue

        # Unordered list
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            indent_spaces = len(m.group(1))
            text = m.group(2)
            level = min(indent_spaces // 2, 3)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
            p.paragraph_format.space_after = Pt(2)
            add_inline_runs(p, text, base_size=Pt(10.5))
            i += 1
            continue

        # Ordered list
        m = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m:
            indent_spaces = len(m.group(1))
            text = m.group(3)
            level = min(indent_spaces // 2, 3)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
            p.paragraph_format.space_after = Pt(2)
            add_inline_runs(p, text, base_size=Pt(10.5))
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        add_inline_runs(p, stripped, base_size=Pt(10.5))
        i += 1


def setup_document(doc: Document) -> None:
    """Configure the default document styles, margins, and font."""
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # Heading styles
    for lvl, size in [(1, 22), (2, 16), (3, 13), (4, 11), (5, 11), (6, 11)]:
        try:
            hstyle = doc.styles[f"Heading {lvl}"]
            hstyle.font.name = "Calibri"
            hstyle.font.size = Pt(size)
            hstyle.font.bold = True
        except KeyError:
            pass


def main() -> int:
    if not SRC.exists():
        print(f"ERROR: source not found: {SRC}", file=sys.stderr)
        return 1
    md_text = SRC.read_text(encoding="utf-8")
    doc = Document()
    setup_document(doc)
    convert(md_text, doc)
    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DST))
    print(f"Wrote {DST} ({DST.stat().st_size:,} bytes)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())